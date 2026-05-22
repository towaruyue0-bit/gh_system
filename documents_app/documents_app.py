# =============================================================================
# 書類差し込みアプリ
# グループホーム 利用者書類管理システム
# =============================================================================
#
# 【使い方】
#   1. documents_app/templates/ フォルダに Excel(.xlsx) か Word(.docx) の
#      テンプレートを入れておく
#   2. テンプレート内で差し込みたい箇所に {{氏名}} などのプレースホルダーを書く
#   3. このアプリで利用者と書類を選んで「差し込んで開く」をクリック
#
# 【必要なライブラリ】
#   コマンドプロンプトで以下を実行してください：
#   pip install openpyxl python-docx
#
# =============================================================================

import tkinter as tk
from tkinter import ttk, messagebox
import sqlite3
import os
import glob
from datetime import datetime, date

# 外部ライブラリの確認（pip install openpyxl python-docx が必要）
try:
    import openpyxl
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# =============================================================================
# パス設定
# =============================================================================
BASE_DIR = r"C:\Users\Public\gh_system"
PRIVATE_ROOT = r"C:\GH_Data"
DATA_DIR     = os.path.join(PRIVATE_ROOT, "data")
DB_PATH      = os.path.join(DATA_DIR, "residents.db")
APP_DIR      = os.path.join(BASE_DIR, "documents_app")
TEMPLATE_DIR = os.path.join(APP_DIR, "templates")
TEMP_DIR     = os.path.join(APP_DIR, "temp")

# フォント設定（日本語対応）
FONT       = ("MS Gothic", 11)
FONT_BOLD  = ("MS Gothic", 11, "bold")
FONT_TITLE = ("MS Gothic", 14, "bold")
FONT_SMALL = ("MS Gothic", 10)

# カラー設定
COLOR_BG          = "#F5F7FA"
COLOR_HEADER      = "#2C5F8A"
COLOR_BTN_PRIMARY = "#2C5F8A"
COLOR_BTN_OPEN    = "#27AE60"
COLOR_BTN_HELP    = "#7F8C8D"
COLOR_ROW_ODD     = "#FFFFFF"
COLOR_ROW_EVEN    = "#F0F4FA"

# =============================================================================
# 利用できるプレースホルダー一覧
# テンプレートファイルにこの文字列を書いておくと自動で置き換わる
# =============================================================================
PLACEHOLDER_HELP = """\
テンプレートファイル内に以下の文字列を書くと、利用者情報に自動置換されます。
文字列はそのまま（{{ }} を含めて）コピーして使ってください。

【利用者情報】
  {{氏名}}             → 利用者の氏名
  {{ふりがな}}         → ふりがな
  {{部屋番号}}         → 部屋番号
  {{生年月日}}         → 生年月日（例: 1980年4月1日）
  {{年齢}}             → 現在の年齢（数字のみ）
  {{入居日}}           → 入居日（例: 2020年4月1日）
  {{障害支援区分}}     → 障害支援区分（数字）
  {{受給者番号}}       → 受給者番号
  {{緊急連絡先氏名}}   → 緊急連絡先の氏名
  {{緊急連絡先電話}}   → 緊急連絡先の電話番号
  {{緊急連絡先続柄}}   → 続柄（例: 長女）

【今日の日付】
  {{今日の日付}}       → 本日の日付（例: 2025年6月1日）
  {{今日の年}}         → 年のみ（例: 2025）
  {{今日の月}}         → 月のみ（例: 6）
  {{今日の日}}         → 日のみ（例: 1）

【注意事項】
  ・Wordの場合、プレースホルダーは途中で書式を変えずに
    1つのまとまり（ラン）として入力してください。
  ・Excelの場合、セルに直接テキストとして入力してください。
  ・差し込み後のファイルは temp フォルダに一時保存されます。
    印刷が終わったら保存する場合は別の場所に移動してください。
"""

# =============================================================================
# セットアップ関数
# =============================================================================

def setup_directories():
    """必要なフォルダを自動作成する"""
    for d in [TEMPLATE_DIR, TEMP_DIR]:
        os.makedirs(d, exist_ok=True)


def cleanup_temp():
    """起動時に古い一時ファイルを削除する"""
    if not os.path.exists(TEMP_DIR):
        return
    for f in glob.glob(os.path.join(TEMP_DIR, "*")):
        try:
            os.remove(f)
        except Exception:
            pass  # 開いているファイルは削除できないので無視する

# =============================================================================
# DB操作
# =============================================================================

def get_conn():
    """SQLiteへの接続を返す（列名でアクセスできるよう Row ファクトリを設定）"""
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn


def fetch_active_residents():
    """入居中の利用者を部屋番号・氏名順で取得する"""
    try:
        conn = get_conn()
        rows = conn.execute(
            "SELECT * FROM residents WHERE status = '入居中' ORDER BY room_number, name"
        ).fetchall()
        conn.close()
        return rows
    except Exception:
        return []

# =============================================================================
# 差し込み処理
# =============================================================================

def calc_age(birthdate_str):
    """
    生年月日の文字列（YYYY-MM-DD）から現在の年齢を計算して返す。
    計算できない場合は空文字を返す。
    """
    if not birthdate_str:
        return ""
    try:
        bd    = date.fromisoformat(birthdate_str.strip())
        today = date.today()
        age   = today.year - bd.year - ((today.month, today.day) < (bd.month, bd.day))
        return str(age)
    except ValueError:
        return ""


def format_date(date_str):
    """
    YYYY-MM-DD 形式の日付を「YYYY年M月D日」形式に変換する。
    変換できない場合はそのまま返す。
    """
    if not date_str:
        return ""
    try:
        d = date.fromisoformat(date_str.strip())
        return f"{d.year}年{d.month}月{d.day}日"
    except ValueError:
        return date_str


def build_replace_map(resident):
    """
    利用者データからプレースホルダー置換マップを作成する。
    戻り値: {プレースホルダー文字列: 置換後の値} の辞書
    """
    today = date.today()
    return {
        "{{氏名}}":           resident["name"]                       or "",
        "{{ふりがな}}":       resident["furigana"]                   or "",
        "{{部屋番号}}":       resident["room_number"]                or "",
        "{{生年月日}}":       format_date(resident["birthdate"]),
        "{{年齢}}":           calc_age(resident["birthdate"]),
        "{{入居日}}":         format_date(resident["move_in_date"]),
        "{{障害支援区分}}":   str(resident["disability_grade"])      if resident["disability_grade"] else "",
        "{{受給者番号}}":     resident["recipient_number"]           or "",
        "{{緊急連絡先氏名}}": resident["emergency_contact_name"]     or "",
        "{{緊急連絡先電話}}": resident["emergency_contact_phone"]    or "",
        "{{緊急連絡先続柄}}": resident["emergency_contact_relation"] or "",
        "{{今日の日付}}":     f"{today.year}年{today.month}月{today.day}日",
        "{{今日の年}}":       str(today.year),
        "{{今日の月}}":       str(today.month),
        "{{今日の日}}":       str(today.day),
    }


def replace_in_xlsx(template_path, replace_map, output_path):
    """
    Excelテンプレートのプレースホルダーを置換して output_path に保存する。
    全シートの全セルを対象にテキスト置換を行う。
    """
    wb = openpyxl.load_workbook(template_path)
    for sheet in wb.worksheets:
        for row in sheet.iter_rows():
            for cell in row:
                if cell.value and isinstance(cell.value, str):
                    for key, val in replace_map.items():
                        if key in cell.value:
                            cell.value = cell.value.replace(key, val)
    wb.save(output_path)


def replace_in_docx(template_path, replace_map, output_path):
    """
    Wordテンプレートのプレースホルダーを置換して output_path に保存する。
    段落・表のセルを対象にテキスト置換を行う。

    注意: プレースホルダーは書式を変えずに1つのラン（Run）として書いてください。
    途中で太字などを変えると分割されて置換できないことがあります。
    """
    doc = DocxDocument(template_path)

    def _replace_in_paragraph(para):
        """段落内の各ラン（書式の固まり）でプレースホルダーを置換する"""
        for run in para.runs:
            for key, val in replace_map.items():
                if key in run.text:
                    run.text = run.text.replace(key, val)

    # 通常の段落
    for para in doc.paragraphs:
        _replace_in_paragraph(para)

    # 表の中のセル
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_paragraph(para)

    doc.save(output_path)


def create_merged_file(template_path, resident):
    """
    テンプレートに利用者データを差し込んだ一時ファイルを作成する。
    成功したらファイルパスを返す。失敗したら None を返す。
    """
    ext           = os.path.splitext(template_path)[1].lower()
    name          = resident["name"] or "利用者"
    timestamp     = datetime.now().strftime("%Y%m%d_%H%M%S_%f")
    template_name = os.path.splitext(os.path.basename(template_path))[0]
    output_name   = f"{name}_{template_name}_{timestamp}{ext}"
    output_path   = os.path.join(TEMP_DIR, output_name)

    replace_map = build_replace_map(resident)

    try:
        if ext == ".xlsx":
            if not OPENPYXL_AVAILABLE:
                raise ImportError(
                    "openpyxl がインストールされていません。\n"
                    "コマンドプロンプトで「pip install openpyxl」を実行してください。"
                )
            replace_in_xlsx(template_path, replace_map, output_path)

        elif ext == ".docx":
            if not DOCX_AVAILABLE:
                raise ImportError(
                    "python-docx がインストールされていません。\n"
                    "コマンドプロンプトで「pip install python-docx」を実行してください。"
                )
            replace_in_docx(template_path, replace_map, output_path)

        else:
            raise ValueError(f"対応していないファイル形式です: {ext}\n（対応形式: .xlsx / .docx）")

        return output_path

    except Exception as e:
        messagebox.showerror("エラー", f"ファイルの作成中にエラーが発生しました。\n\n{e}")
        return None

# =============================================================================
# メインアプリ
# =============================================================================

class DocumentsApp(tk.Tk):
    """書類差し込みアプリのメインウィンドウ"""

    def __init__(self):
        super().__init__()
        self.title("書類差し込みアプリ")
        self.geometry("920x580")
        self.minsize(700, 460)
        self.configure(bg=COLOR_BG)

        # 内部データ（利用者・テンプレートの一覧）
        self.residents = []
        self.templates = []

        self._check_libraries()
        self._build_ui()
        self._load_residents()
        self._load_templates()

    # ------------------------------------------------------------------
    # 起動チェック
    # ------------------------------------------------------------------

    def _check_libraries(self):
        """必要ライブラリが不足している場合に案内ダイアログを表示する"""
        missing = []
        if not OPENPYXL_AVAILABLE:
            missing.append("openpyxl   （Excel用）")
        if not DOCX_AVAILABLE:
            missing.append("python-docx（Word用）")
        if missing:
            msg = (
                "以下のライブラリがインストールされていません。\n\n"
                + "\n".join(f"  ・{m}" for m in missing)
                + "\n\n【インストール方法】\n"
                "コマンドプロンプト（黒い画面）で以下を実行してください：\n\n"
                "  pip install openpyxl python-docx\n\n"
                "インストール後にアプリを再起動してください。"
            )
            messagebox.showwarning("ライブラリ未インストール", msg)

    # ------------------------------------------------------------------
    # UI構築
    # ------------------------------------------------------------------

    def _build_ui(self):
        """画面全体のレイアウトを構築する"""

        # ── ヘッダー ──────────────────────────────────────────────────
        header = tk.Frame(self, bg=COLOR_HEADER, height=52)
        header.pack(fill="x")
        header.pack_propagate(False)

        tk.Label(
            header,
            text="書類差し込みアプリ",
            font=FONT_TITLE,
            fg="white",
            bg=COLOR_HEADER,
        ).pack(side="left", padx=20, pady=12)

        tk.Button(
            header,
            text="？ プレースホルダー一覧",
            font=FONT_SMALL,
            bg="#4A7FA8",
            fg="white",
            relief="flat",
            cursor="hand2",
            padx=10,
            pady=4,
            command=self._show_placeholder_help,
        ).pack(side="right", padx=15, pady=12)

        # ── メインエリア（左右2パネル） ─────────────────────────────
        main = tk.Frame(self, bg=COLOR_BG)
        main.pack(fill="both", expand=True, padx=15, pady=10)

        # 左パネル：利用者一覧
        left = tk.LabelFrame(
            main,
            text="① 利用者を選ぶ（複数選択可）",
            font=FONT_BOLD,
            bg=COLOR_BG,
            padx=6,
            pady=6,
        )
        left.pack(side="left", fill="both", expand=True, padx=(0, 8))

        # リストボックス＋スクロールバーをまとめるフレーム
        r_frame = tk.Frame(left, bd=1, relief="solid")
        r_frame.pack(fill="both", expand=True)

        r_scroll = ttk.Scrollbar(r_frame, orient="vertical")
        r_scroll.pack(side="right", fill="y")

        self.resident_lb = tk.Listbox(
            r_frame,
            font=FONT,
            selectmode="extended",   # Ctrl/Shift で複数選択できる
            activestyle="none",
            bd=0,
            relief="flat",
            yscrollcommand=r_scroll.set,
        )
        self.resident_lb.pack(side="left", fill="both", expand=True)
        r_scroll.configure(command=self.resident_lb.yview)
        self.resident_lb.bind("<<ListboxSelect>>", self._on_select)

        tk.Button(
            left,
            text="全員選択",
            font=FONT_SMALL,
            bg="#E0E8F0",
            relief="flat",
            cursor="hand2",
            pady=4,
            command=self._select_all,
        ).pack(fill="x", pady=(6, 0))

        # 右パネル：テンプレート一覧
        right = tk.LabelFrame(
            main,
            text="② 書類テンプレートを選ぶ",
            font=FONT_BOLD,
            bg=COLOR_BG,
            padx=6,
            pady=6,
        )
        right.pack(side="right", fill="both", expand=True, padx=(8, 0))

        t_frame = tk.Frame(right, bd=1, relief="solid")
        t_frame.pack(fill="both", expand=True)

        t_scroll = ttk.Scrollbar(t_frame, orient="vertical")
        t_scroll.pack(side="right", fill="y")

        self.template_lb = tk.Listbox(
            t_frame,
            font=FONT,
            selectmode="single",
            activestyle="none",
            bd=0,
            relief="flat",
            yscrollcommand=t_scroll.set,
        )
        self.template_lb.pack(side="left", fill="both", expand=True)
        t_scroll.configure(command=self.template_lb.yview)
        self.template_lb.bind("<<ListboxSelect>>", self._on_select)

        tk.Button(
            right,
            text="テンプレートフォルダを開く",
            font=FONT_SMALL,
            bg="#E0E8F0",
            relief="flat",
            cursor="hand2",
            pady=4,
            command=self._open_folder,
        ).pack(fill="x", pady=(6, 0))

        # ── 下部ボタンエリア ──────────────────────────────────────────
        bottom = tk.Frame(self, bg=COLOR_BG)
        bottom.pack(fill="x", padx=15, pady=(0, 14))

        # ステータス表示
        self.status_lbl = tk.Label(
            bottom,
            text="利用者と書類を選んでください",
            font=FONT_SMALL,
            fg="#666",
            bg=COLOR_BG,
        )
        self.status_lbl.pack(side="left")

        # 差し込みボタン（右端）
        self.merge_btn = tk.Button(
            bottom,
            text="▶ 差し込んで開く",
            font=FONT_BOLD,
            bg=COLOR_BTN_OPEN,
            fg="white",
            relief="flat",
            cursor="hand2",
            padx=20,
            pady=8,
            state="disabled",
            command=self._do_merge,
        )
        self.merge_btn.pack(side="right")

        # 再読み込みボタン
        tk.Button(
            bottom,
            text="↻ 再読み込み",
            font=FONT_SMALL,
            bg="#E0E8F0",
            relief="flat",
            cursor="hand2",
            padx=12,
            pady=8,
            command=self._reload,
        ).pack(side="right", padx=(0, 10))

    # ------------------------------------------------------------------
    # データ読み込み
    # ------------------------------------------------------------------

    def _load_residents(self):
        """入居中の利用者をDBから読み込んでリストに表示する"""
        self.residents = fetch_active_residents()
        self.resident_lb.delete(0, "end")

        if not self.residents:
            self.resident_lb.insert("end", "（入居中の利用者がいません）")
            return

        for i, r in enumerate(self.residents):
            room  = f"[{r['room_number']}]" if r["room_number"] else "[--]"
            label = f"{room}  {r['name']}"
            self.resident_lb.insert("end", label)
            color = COLOR_ROW_ODD if i % 2 == 0 else COLOR_ROW_EVEN
            self.resident_lb.itemconfig(i, bg=color)

    def _load_templates(self):
        """テンプレートフォルダから .xlsx / .docx ファイルを読み込んでリストに表示する"""
        self.templates = sorted(
            glob.glob(os.path.join(TEMPLATE_DIR, "*.xlsx"))
            + glob.glob(os.path.join(TEMPLATE_DIR, "*.docx"))
        )
        self.template_lb.delete(0, "end")

        if not self.templates:
            self.template_lb.insert("end", "（テンプレートがありません）")
            self.template_lb.insert("end", "↑「テンプレートフォルダを開く」から追加")
            return

        for i, path in enumerate(self.templates):
            basename = os.path.basename(path)
            name, ext = os.path.splitext(basename)
            label = f"[{ext[1:].upper()}]  {name}"
            self.template_lb.insert("end", label)
            color = COLOR_ROW_ODD if i % 2 == 0 else COLOR_ROW_EVEN
            self.template_lb.itemconfig(i, bg=color)

    # ------------------------------------------------------------------
    # イベントハンドラ
    # ------------------------------------------------------------------

    def _on_select(self, event=None):
        """利用者またはテンプレートの選択が変わったときにボタンとステータスを更新する"""
        r_sel = self.resident_lb.curselection()
        t_sel = self.template_lb.curselection()

        valid_r = len(r_sel) > 0 and len(self.residents) > 0
        valid_t = len(t_sel) > 0 and len(self.templates) > 0

        if valid_r and valid_t:
            t_name = os.path.splitext(os.path.basename(self.templates[t_sel[0]]))[0]
            self.status_lbl.config(
                text=f"{len(r_sel)}名 ×「{t_name}」を差し込みます",
                fg=COLOR_BTN_PRIMARY,
            )
            self.merge_btn.config(state="normal")
        else:
            self.status_lbl.config(text="利用者と書類を選んでください", fg="#666")
            self.merge_btn.config(state="disabled")

    def _select_all(self):
        """利用者リストを全員選択する"""
        if self.residents:
            self.resident_lb.selection_set(0, "end")
            self._on_select()

    def _open_folder(self):
        """テンプレートフォルダをエクスプローラーで開く"""
        os.startfile(TEMPLATE_DIR)

    def _reload(self):
        """利用者とテンプレートを再読み込みする"""
        self._load_residents()
        self._load_templates()
        self._on_select()

    def _show_placeholder_help(self):
        """プレースホルダー一覧をダイアログで表示する"""
        dialog = tk.Toplevel(self)
        dialog.title("プレースホルダー一覧")
        dialog.geometry("560x520")
        dialog.resizable(False, True)
        dialog.grab_set()

        # スクロール可能なテキストエリア
        frame = tk.Frame(dialog, padx=12, pady=10)
        frame.pack(fill="both", expand=True)

        sb = ttk.Scrollbar(frame, orient="vertical")
        sb.pack(side="right", fill="y")

        txt = tk.Text(
            frame,
            font=FONT_SMALL,
            wrap="word",
            bd=0,
            relief="flat",
            bg="#F9F9F9",
            yscrollcommand=sb.set,
            state="normal",
        )
        txt.pack(side="left", fill="both", expand=True)
        sb.configure(command=txt.yview)

        txt.insert("1.0", PLACEHOLDER_HELP)
        txt.config(state="disabled")  # 読み取り専用にする

        tk.Button(
            dialog,
            text="閉じる",
            font=FONT,
            relief="flat",
            cursor="hand2",
            bg="#E0E8F0",
            padx=20,
            pady=6,
            command=dialog.destroy,
        ).pack(pady=(0, 10))

    # ------------------------------------------------------------------
    # 差し込み実行
    # ------------------------------------------------------------------

    def _do_merge(self):
        """選択された利用者とテンプレートで差し込みを実行してファイルを開く"""
        r_sel = self.resident_lb.curselection()
        t_sel = self.template_lb.curselection()

        if not r_sel or not t_sel:
            return

        template_path      = self.templates[t_sel[0]]
        selected_residents = [self.residents[i] for i in r_sel]

        # 複数名の場合は確認ダイアログを表示する
        if len(selected_residents) > 1:
            ok = messagebox.askyesno(
                "確認",
                f"{len(selected_residents)}名分のファイルを作成して開きます。\n"
                "よろしいですか？\n\n"
                "（ファイルはそれぞれ別ウィンドウで開きます）",
            )
            if not ok:
                return

        # 差し込みを実行してファイルを開く
        opened = 0
        for resident in selected_residents:
            output_path = create_merged_file(template_path, resident)
            if output_path and os.path.exists(output_path):
                os.startfile(output_path)
                opened += 1

        if opened > 0:
            messagebox.showinfo(
                "完了",
                f"{opened}名分のファイルを開きました。\n\n"
                "印刷はファイルが開いてから行ってください。\n"
                "保存する場合は「名前を付けて保存」で別の場所に保存してください。",
            )

# =============================================================================
# 起動
# =============================================================================

if __name__ == "__main__":
    setup_directories()
    cleanup_temp()
    app = DocumentsApp()
    app.mainloop()
